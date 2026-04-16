import io
import json
import logging
import warnings
from pathlib import Path

import anthropic
from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger(__name__)

DOCS_FOLDER = Path(__file__).parent.parent / "docs"


def read_pdf(file_path_or_buffer) -> str:
    try:
        import pdfplumber
        if hasattr(file_path_or_buffer, "read"):
            data = file_path_or_buffer.read()
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        else:
            with pdfplumber.open(file_path_or_buffer) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception as e:
        logger.warning(f"PDF read error: {e}")
        return ""


def read_docx(file_path_or_buffer) -> str:
    try:
        from docx import Document
        if hasattr(file_path_or_buffer, "read"):
            data = file_path_or_buffer.read()
            doc = Document(io.BytesIO(data))
        else:
            doc = Document(file_path_or_buffer)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        logger.warning(f"DOCX read error: {e}")
        return ""


def read_txt(file_path_or_buffer) -> str:
    try:
        if hasattr(file_path_or_buffer, "read"):
            data = file_path_or_buffer.read()
            if isinstance(data, bytes):
                return data.decode("utf-8", errors="replace")
            return data
        else:
            return Path(file_path_or_buffer).read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        logger.warning(f"TXT read error: {e}")
        return ""


def read_any_file(file_path_or_buffer, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return read_pdf(file_path_or_buffer)
    elif ext == ".docx":
        return read_docx(file_path_or_buffer)
    elif ext == ".txt":
        return read_txt(file_path_or_buffer)
    else:
        logger.warning(f"Unsupported file format: {filename}")
        return ""


def load_from_docs_folder() -> dict:
    if not DOCS_FOLDER.exists():
        return {"files_found": [], "raw_texts": {}, "combined_text": ""}

    files_found = []
    raw_texts = {}
    parts = []

    for path in sorted(DOCS_FOLDER.iterdir()):
        if path.suffix.lower() in (".pdf", ".docx", ".txt") and path.is_file():
            text = read_any_file(path, path.name)
            if text.strip():
                files_found.append(path.name)
                raw_texts[path.name] = text
                parts.append(f"=== {path.name} ===\n{text}\n")

    combined_text = "\n".join(parts)
    return {"files_found": files_found, "raw_texts": raw_texts, "combined_text": combined_text}


def load_from_uploads(uploaded_files: list) -> dict:
    files_found = []
    raw_texts = {}
    parts = []

    for uf in uploaded_files:
        filename = uf.name
        text = read_any_file(uf, filename)
        if text.strip():
            files_found.append(filename)
            raw_texts[filename] = text
            parts.append(f"=== {filename} ===\n{text}\n")

    combined_text = "\n".join(parts)
    return {"files_found": files_found, "raw_texts": raw_texts, "combined_text": combined_text}


def parse_context_with_claude(combined_text: str) -> dict:
    client = anthropic.Anthropic()

    system_prompt = (
        "You are a personal career context parser. The user has provided one or more career documents. "
        "Extract and structure everything into a single unified JSON object with these exact keys:\n\n"
        "name (string),\n"
        "phone (string),\n"
        "email (string),\n"
        "linkedin (string),\n"
        "location (string),\n"
        "summary (string — best professional summary found),\n"
        "skills (list of strings),\n"
        "experience (list of objects: title, company, dates, bullets as list of strings),\n"
        "projects (list of objects: name, tech, github, bullets as list of strings),\n"
        "education (list of objects: degree, school, dates),\n"
        "work_stories (list of objects: situation, task, action, result, theme — extract STAR stories if present),\n"
        "about_me (string — elevator pitch if found, else null),\n"
        "bio (string — LinkedIn bio or narrative if found, else null),\n"
        "additional_context (string — any other relevant career info)\n\n"
        "If a field is not found set it to null.\n"
        "Return only valid JSON. No preamble, no markdown, no backticks."
    )

    try:
        message = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=8096,
            system=system_prompt,
            messages=[{"role": "user", "content": combined_text}],
        )
        raw = message.content[0].text.strip()
        # Strip markdown fences if present
        if raw.startswith("```"):
            parts = raw.split("```")
            raw = parts[1] if len(parts) > 1 else raw
            if raw.startswith("json"):
                raw = raw[4:]
        raw = raw.strip()
        return json.loads(raw)
    except json.JSONDecodeError as e:
        logger.error(f"Context JSON parse error: {e} | raw snippet: {raw[:300] if 'raw' in dir() else 'N/A'}")
        return {"raw_text": combined_text, "parse_error": True, "error_detail": str(e)}
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Context parse error: {type(e).__name__}: {error_msg}")
        # Surface billing/auth errors clearly
        if "credit balance" in error_msg or "billing" in error_msg.lower():
            return {"raw_text": combined_text, "parse_error": True, "error_detail": "Anthropic API credit balance too low. Go to console.anthropic.com → Plans & Billing to add credits."}
        if "authentication" in error_msg.lower() or "api_key" in error_msg.lower():
            return {"raw_text": combined_text, "parse_error": True, "error_detail": "Invalid ANTHROPIC_API_KEY. Check your .env file."}
        return {"raw_text": combined_text, "parse_error": True, "error_detail": error_msg}


def build_context(uploaded_files: list = None) -> dict:
    if uploaded_files:
        loaded = load_from_uploads(uploaded_files)
    else:
        loaded = load_from_docs_folder()

    if not loaded["combined_text"].strip():
        return {"empty": True}

    parsed = parse_context_with_claude(loaded["combined_text"])
    parsed["files_found"] = loaded["files_found"]
    return parsed
